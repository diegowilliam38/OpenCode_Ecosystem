#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gera template DOCX com formatação ABNT para uso como reference-doc do pandoc.
Uso: python gerar_template_abnt.py [--output template-abnt.docx]
"""

import argparse
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def criar_template_abnt(output_path="template-abnt.docx"):
    """Cria documento template ABNT formatado programaticamente."""
    doc = Document()

    # ─── Margens ABNT ───
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # ─── Estilo Normal ───
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ─── Estilos de Título (ABNT NBR 6024) ───
    config_titulos = {
        'Heading 1':  ('Heading 1', 14, True,  Pt(24), Pt(12)),
        'Heading 2':  ('Heading 2', 13, True,  Pt(18), Pt(6)),
        'Heading 3':  ('Heading 3', 12, True,  Pt(12), Pt(3)),
        'Title':      ('Title',      16, True,  Pt(36), Pt(18)),
        'Subtitle':   ('Subtitle',   14, False, Pt(12), Pt(6)),
    }
    for nome, (_, sz, bold, antes, depois) in config_titulos.items():
        if nome in doc.styles:
            s = doc.styles[nome]
            s.font.name = 'Times New Roman'
            s.font.size = Pt(sz)
            s.font.bold = bold
            s.font.color.rgb = RGBColor(0, 0, 0)
            s.paragraph_format.space_before = antes
            s.paragraph_format.space_after = depois
            s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ─── Estilo Citação (ABNT NBR 10520) ───
    if 'Quote' in doc.styles:
        q = doc.styles['Quote']
        q.font.name = 'Times New Roman'
        q.font.size = Pt(10)
        q.font.italic = True
        q.paragraph_format.left_indent = Cm(4)
        q.paragraph_format.line_spacing = 1.0
        q.paragraph_format.space_before = Pt(6)
        q.paragraph_format.space_after = Pt(6)

    # ─── Numeração automática de páginas ───
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)

    # ─── Salvar ───
    doc.save(output_path)
    print(f"Template ABNT salvo em: {output_path}")
    return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Gera template DOCX ABNT")
    parser.add_argument("--output", default="template-abnt.docx", help="Caminho de saída")
    args = parser.parse_args()
    criar_template_abnt(args.output)
